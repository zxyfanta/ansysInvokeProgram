"""
数据分析 - 报告生成器

生成专业的分析报告。
"""

import os
import logging
from typing import Dict, List, Any, Optional
from pathlib import Path
from datetime import datetime
import json

# 报告模板导入
try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image, Table, TableStyle, PageBreak
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from reportlab.pdfgen import canvas
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False
    print("ReportLab不可用，PDF报告功能将受限")

# Word文档生成
try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("python-docx不可用，Word报告功能将受限")

class ReportGenerator:
    """报告生成器"""
    
    def __init__(self):
        self.logger = logging.getLogger(__name__)
        
        # 报告配置
        self.report_config = {
            'title': '激光毁伤效能分析报告',
            'author': '激光毁伤效能分析软件',
            'organization': '军用软件开发部门',
            'classification': '内部资料',
            'template_style': 'professional'
        }
        
        # 样式配置
        if REPORTLAB_AVAILABLE:
            self.styles = getSampleStyleSheet()
            self._setup_custom_styles()
    
    def _setup_custom_styles(self):
        """设置自定义样式"""
        if not REPORTLAB_AVAILABLE:
            return
        
        # 标题样式
        self.title_style = ParagraphStyle(
            'CustomTitle',
            parent=self.styles['Title'],
            fontSize=24,
            spaceAfter=30,
            alignment=1,  # 居中
            textColor=colors.darkblue
        )
        
        # 章节标题样式
        self.heading_style = ParagraphStyle(
            'CustomHeading',
            parent=self.styles['Heading1'],
            fontSize=16,
            spaceAfter=12,
            textColor=colors.darkblue,
            borderWidth=1,
            borderColor=colors.darkblue,
            borderPadding=5
        )
        
        # 正文样式
        self.body_style = ParagraphStyle(
            'CustomBody',
            parent=self.styles['Normal'],
            fontSize=12,
            spaceAfter=6,
            alignment=0  # 左对齐
        )
    
    def generate_comprehensive_report(self, analysis_data: Dict, chart_files: List[str], 
                                    output_file: str, format: str = 'pdf') -> bool:
        """生成综合分析报告"""
        try:
            if format.lower() == 'pdf' and REPORTLAB_AVAILABLE:
                return self._generate_pdf_report(analysis_data, chart_files, output_file)
            elif format.lower() == 'docx' and DOCX_AVAILABLE:
                return self._generate_word_report(analysis_data, chart_files, output_file)
            elif format.lower() == 'html':
                return self._generate_html_report(analysis_data, chart_files, output_file)
            else:
                self.logger.error(f"不支持的报告格式: {format}")
                return False
                
        except Exception as e:
            self.logger.error(f"报告生成失败: {e}")
            return False
    
    def _generate_pdf_report(self, analysis_data: Dict, chart_files: List[str], output_file: str) -> bool:
        """生成PDF报告"""
        try:
            doc = SimpleDocTemplate(output_file, pagesize=A4)
            story = []
            
            # 报告标题页
            story.extend(self._create_title_page())
            story.append(PageBreak())
            
            # 目录
            story.extend(self._create_table_of_contents())
            story.append(PageBreak())
            
            # 执行摘要
            story.extend(self._create_executive_summary(analysis_data))
            story.append(PageBreak())
            
            # 激光毁伤分析章节
            if 'laser_damage' in analysis_data:
                story.extend(self._create_laser_damage_section(analysis_data['laser_damage'], chart_files))
                story.append(PageBreak())
            
            # 毁伤后效分析章节
            if 'post_damage' in analysis_data:
                story.extend(self._create_post_damage_section(analysis_data['post_damage'], chart_files))
                story.append(PageBreak())
            
            # 对比分析章节
            if 'comparison' in analysis_data:
                story.extend(self._create_comparison_section(analysis_data['comparison'], chart_files))
                story.append(PageBreak())
            
            # 结论和建议
            story.extend(self._create_conclusions_section(analysis_data))
            
            # 附录
            story.extend(self._create_appendix_section(analysis_data))
            
            # 生成PDF
            doc.build(story)
            
            self.logger.info(f"PDF报告已生成: {output_file}")
            return True
            
        except Exception as e:
            self.logger.error(f"PDF报告生成失败: {e}")
            return False
    
    def _generate_word_report(self, analysis_data: Dict, chart_files: List[str], output_file: str) -> bool:
        """生成Word报告"""
        try:
            doc = Document()
            
            # 设置文档样式
            self._setup_word_styles(doc)
            
            # 标题页
            self._add_word_title_page(doc)
            doc.add_page_break()
            
            # 执行摘要
            self._add_word_executive_summary(doc, analysis_data)
            doc.add_page_break()
            
            # 激光毁伤分析
            if 'laser_damage' in analysis_data:
                self._add_word_laser_damage_section(doc, analysis_data['laser_damage'], chart_files)
                doc.add_page_break()
            
            # 毁伤后效分析
            if 'post_damage' in analysis_data:
                self._add_word_post_damage_section(doc, analysis_data['post_damage'], chart_files)
                doc.add_page_break()
            
            # 对比分析
            if 'comparison' in analysis_data:
                self._add_word_comparison_section(doc, analysis_data['comparison'], chart_files)
                doc.add_page_break()
            
            # 结论
            self._add_word_conclusions_section(doc, analysis_data)
            
            # 保存文档
            doc.save(output_file)
            
            self.logger.info(f"Word报告已生成: {output_file}")
            return True
            
        except Exception as e:
            self.logger.error(f"Word报告生成失败: {e}")
            return False
    
    def _generate_html_report(self, analysis_data: Dict, chart_files: List[str], output_file: str) -> bool:
        """生成HTML报告"""
        try:
            html_content = self._create_html_template(analysis_data, chart_files)
            
            with open(output_file, 'w', encoding='utf-8') as f:
                f.write(html_content)
            
            self.logger.info(f"HTML报告已生成: {output_file}")
            return True
            
        except Exception as e:
            self.logger.error(f"HTML报告生成失败: {e}")
            return False
    
    def _create_title_page(self) -> List:
        """创建标题页"""
        story = []
        
        # 标题
        title = Paragraph(self.report_config['title'], self.title_style)
        story.append(title)
        story.append(Spacer(1, 0.5*inch))
        
        # 副标题
        subtitle = Paragraph("基于ANSYS 2021 R1的激光毁伤效能分析", self.body_style)
        story.append(subtitle)
        story.append(Spacer(1, 1*inch))
        
        # 报告信息表格
        report_info = [
            ['报告类型', '激光毁伤效能分析报告'],
            ['生成时间', datetime.now().strftime('%Y年%m月%d日 %H:%M:%S')],
            ['生成软件', self.report_config['author']],
            ['机构', self.report_config['organization']],
            ['密级', self.report_config['classification']]
        ]
        
        table = Table(report_info, colWidths=[2*inch, 4*inch])
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.lightblue),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        
        story.append(table)
        
        return story
    
    def _create_table_of_contents(self) -> List:
        """创建目录"""
        story = []
        
        story.append(Paragraph("目录", self.heading_style))
        story.append(Spacer(1, 0.2*inch))
        
        toc_items = [
            "1. 执行摘要",
            "2. 激光毁伤分析",
            "   2.1 毁伤指标分析",
            "   2.2 温度场分析",
            "   2.3 应力场分析",
            "3. 毁伤后效分析",
            "   3.1 气动力特性变化",
            "   3.2 飞行轨迹分析",
            "   3.3 稳定性分析",
            "4. 对比分析",
            "   4.1 性能对比",
            "   4.2 轨迹偏差分析",
            "5. 结论和建议",
            "附录A: 技术参数",
            "附录B: 计算方法"
        ]
        
        for item in toc_items:
            story.append(Paragraph(item, self.body_style))
            story.append(Spacer(1, 0.1*inch))
        
        return story
    
    def _create_executive_summary(self, analysis_data: Dict) -> List:
        """创建执行摘要"""
        story = []
        
        story.append(Paragraph("执行摘要", self.heading_style))
        story.append(Spacer(1, 0.2*inch))
        
        # 摘要内容
        summary_text = """
        本报告基于ANSYS 2021 R1平台，采用PyANSYS接口，对激光武器的毁伤效能进行了全面分析。
        分析包括激光毁伤仿真、毁伤后效评估和综合效能评价三个主要方面。
        
        主要发现：
        """
        
        story.append(Paragraph(summary_text, self.body_style))
        
        # 关键指标摘要
        if 'summary' in analysis_data:
            summary = analysis_data['summary']
            key_findings = [
                f"• 最大毁伤温度: {summary.get('max_temperature', 'N/A')} K",
                f"• 毁伤体积比例: {summary.get('damage_ratio', 'N/A')}%",
                f"• 飞行性能退化: {summary.get('performance_degradation', 'N/A')}%",
                f"• 总体毁伤效能: {summary.get('overall_effectiveness', 'N/A')}%"
            ]
            
            for finding in key_findings:
                story.append(Paragraph(finding, self.body_style))
        
        return story
    
    def _create_laser_damage_section(self, laser_data: Dict, chart_files: List[str]) -> List:
        """创建激光毁伤分析章节"""
        story = []
        
        story.append(Paragraph("激光毁伤分析", self.heading_style))
        story.append(Spacer(1, 0.2*inch))
        
        # 毁伤指标分析
        if 'damage_metrics' in laser_data:
            story.append(Paragraph("2.1 毁伤指标分析", self.body_style))
            
            metrics = laser_data['damage_metrics']
            metrics_table = [
                ['指标', '数值', '单位'],
                ['最高温度', f"{metrics.get('max_temperature', 0):.1f}", 'K'],
                ['最大应力', f"{metrics.get('max_stress', 0):.0f}", 'Pa'],
                ['毁伤体积', f"{metrics.get('damage_volume', 0):.6f}", 'm³'],
                ['毁伤深度', f"{metrics.get('damage_depth', 0):.3f}", 'm']
            ]
            
            table = Table(metrics_table, colWidths=[2*inch, 1.5*inch, 1*inch])
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.lightblue),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 10),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            
            story.append(table)
            story.append(Spacer(1, 0.2*inch))
        
        # 添加相关图表
        damage_charts = [f for f in chart_files if 'temperature' in f or 'stress' in f or 'damage' in f]
        for chart_file in damage_charts:
            if os.path.exists(chart_file):
                story.append(Image(chart_file, width=6*inch, height=4*inch))
                story.append(Spacer(1, 0.2*inch))
        
        return story
    
    def _create_post_damage_section(self, post_damage_data: Dict, chart_files: List[str]) -> List:
        """创建毁伤后效分析章节"""
        story = []
        
        story.append(Paragraph("毁伤后效分析", self.heading_style))
        story.append(Spacer(1, 0.2*inch))
        
        # 气动力特性变化
        if 'aerodynamic_coefficients' in post_damage_data:
            story.append(Paragraph("3.1 气动力特性变化", self.body_style))
            
            aero_coeffs = post_damage_data['aerodynamic_coefficients']
            aero_table = [['系数', '数值']]
            
            for coeff, value in aero_coeffs.items():
                aero_table.append([coeff, f"{value:.4f}"])
            
            table = Table(aero_table, colWidths=[2*inch, 2*inch])
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.lightblue),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            
            story.append(table)
            story.append(Spacer(1, 0.2*inch))
        
        # 添加轨迹图表
        trajectory_charts = [f for f in chart_files if 'trajectory' in f or 'altitude' in f or 'velocity' in f]
        for chart_file in trajectory_charts:
            if os.path.exists(chart_file):
                story.append(Image(chart_file, width=6*inch, height=4*inch))
                story.append(Spacer(1, 0.2*inch))
        
        return story
    
    def _create_comparison_section(self, comparison_data: Dict, chart_files: List[str]) -> List:
        """创建对比分析章节"""
        story = []
        
        story.append(Paragraph("对比分析", self.heading_style))
        story.append(Spacer(1, 0.2*inch))
        
        # 性能变化分析
        if 'performance_comparison' in comparison_data:
            story.append(Paragraph("4.1 性能变化分析", self.body_style))
            
            perf_comp = comparison_data['performance_comparison']
            if 'change_percentages' in perf_comp:
                change_table = [['性能指标', '变化百分比']]
                
                for metric, change in perf_comp['change_percentages'].items():
                    change_table.append([metric.replace('_', ' ').title(), f"{change:.1f}%"])
                
                table = Table(change_table, colWidths=[3*inch, 2*inch])
                table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.lightblue),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black)
                ]))
                
                story.append(table)
                story.append(Spacer(1, 0.2*inch))
        
        # 添加对比图表
        comparison_charts = [f for f in chart_files if 'comparison' in f or 'degradation' in f]
        for chart_file in comparison_charts:
            if os.path.exists(chart_file):
                story.append(Image(chart_file, width=6*inch, height=4*inch))
                story.append(Spacer(1, 0.2*inch))
        
        return story
    
    def _create_conclusions_section(self, analysis_data: Dict) -> List:
        """创建结论和建议章节"""
        story = []
        
        story.append(Paragraph("结论和建议", self.heading_style))
        story.append(Spacer(1, 0.2*inch))
        
        conclusions_text = """
        基于本次激光毁伤效能分析，得出以下主要结论：
        
        1. 激光毁伤效果评估
        2. 毁伤后效影响分析
        3. 系统性能评价
        4. 改进建议
        """
        
        story.append(Paragraph(conclusions_text, self.body_style))
        
        return story
    
    def _create_appendix_section(self, analysis_data: Dict) -> List:
        """创建附录章节"""
        story = []
        
        story.append(Paragraph("附录", self.heading_style))
        story.append(Spacer(1, 0.2*inch))
        
        # 技术参数附录
        story.append(Paragraph("附录A: 技术参数", self.body_style))
        
        if 'technical_parameters' in analysis_data:
            params = analysis_data['technical_parameters']
            param_text = json.dumps(params, indent=2, ensure_ascii=False)
            story.append(Paragraph(f"<pre>{param_text}</pre>", self.body_style))
        
        return story

    def _setup_word_styles(self, doc):
        """设置Word文档样式"""
        if not DOCX_AVAILABLE:
            return

        # 设置默认字体
        style = doc.styles['Normal']
        style.font.name = '宋体'
        style.font.size = Pt(12)

    def _add_word_title_page(self, doc):
        """添加Word标题页"""
        if not DOCX_AVAILABLE:
            return

        # 标题
        title = doc.add_heading(self.report_config['title'], 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 副标题
        subtitle = doc.add_paragraph('基于ANSYS 2021 R1的激光毁伤效能分析')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _add_word_executive_summary(self, doc, analysis_data: Dict):
        """添加Word执行摘要"""
        if not DOCX_AVAILABLE:
            return

        doc.add_heading('执行摘要', level=1)

        summary_text = """
        本报告基于ANSYS 2021 R1平台，采用PyANSYS接口，对激光武器的毁伤效能进行了全面分析。
        分析包括激光毁伤仿真、毁伤后效评估和综合效能评价三个主要方面。
        """

        doc.add_paragraph(summary_text)

    def _add_word_laser_damage_section(self, doc, laser_data: Dict, chart_files: List[str]):
        """添加Word激光毁伤分析章节"""
        if not DOCX_AVAILABLE:
            return

        doc.add_heading('激光毁伤分析', level=1)

        if 'damage_metrics' in laser_data:
            doc.add_heading('毁伤指标分析', level=2)

    def _add_word_post_damage_section(self, doc, post_damage_data: Dict, chart_files: List[str]):
        """添加Word毁伤后效分析章节"""
        if not DOCX_AVAILABLE:
            return

        doc.add_heading('毁伤后效分析', level=1)

    def _add_word_comparison_section(self, doc, comparison_data: Dict, chart_files: List[str]):
        """添加Word对比分析章节"""
        if not DOCX_AVAILABLE:
            return

        doc.add_heading('对比分析', level=1)

    def _add_word_conclusions_section(self, doc, analysis_data: Dict):
        """添加Word结论章节"""
        if not DOCX_AVAILABLE:
            return

        doc.add_heading('结论和建议', level=1)

        conclusions_text = """
        基于本次激光毁伤效能分析，得出以下主要结论：

        1. 激光毁伤效果评估
        2. 毁伤后效影响分析
        3. 系统性能评价
        4. 改进建议
        """

        doc.add_paragraph(conclusions_text)

    def _create_html_template(self, analysis_data: Dict, chart_files: List[str]) -> str:
        """创建HTML报告模板"""
        html_template = f"""
        <!DOCTYPE html>
        <html lang="zh-CN">
        <head>
            <meta charset="UTF-8">
            <title>{self.report_config['title']}</title>
            <style>
                body {{ font-family: 'Microsoft YaHei', Arial, sans-serif; }}
                .container {{ max-width: 1200px; margin: 0 auto; padding: 30px; }}
                .title {{ color: #2c3e50; font-size: 2.5em; text-align: center; }}
                .section-title {{ color: #2c3e50; font-size: 1.8em; margin: 20px 0; }}
                table {{ width: 100%; border-collapse: collapse; margin: 20px 0; }}
                th, td {{ border: 1px solid #ddd; padding: 12px; }}
                th {{ background-color: #3498db; color: white; }}
            </style>
        </head>
        <body>
            <div class="container">
                <h1 class="title">{self.report_config['title']}</h1>
                <p style="text-align: center;">基于ANSYS 2021 R1的激光毁伤效能分析</p>

                <h2 class="section-title">执行摘要</h2>
                <p>本报告基于ANSYS 2021 R1平台，采用PyANSYS接口，对激光武器的毁伤效能进行了全面分析。</p>

                <p style="text-align: right; color: #7f8c8d;">
                    报告生成时间: {datetime.now().strftime('%Y年%m月%d日 %H:%M:%S')}
                </p>
            </div>
        </body>
        </html>
        """

        return html_template
